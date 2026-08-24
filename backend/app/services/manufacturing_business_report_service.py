import json
from collections import defaultdict
from datetime import datetime, timezone
from io import BytesIO

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.manufacturing import EnergyRecord, ProductionRecord
from app.models.manufacturing_business_report import ManufacturingBusinessReport
from app.schemas.manufacturing_business_report import (
    ManufacturingBusinessReportGenerateRequest,
    ManufacturingBusinessReportSnapshotCreate,
)
from app.services.ai_analysis_service import AIAnalysisService
from app.services.equipment_diagnosis_service import EquipmentDiagnosisService
from app.services.equipment_management_service import EquipmentManagementService


class ManufacturingBusinessReportService:
    """Persistence boundary for immutable manufacturing report snapshots."""

    def build_deterministic_snapshot(self, db: Session) -> dict:
        """Collect persisted manufacturing data and calculate reproducible report facts."""
        production_records = db.scalars(
            select(ProductionRecord).order_by(ProductionRecord.date.asc(), ProductionRecord.id.asc())
        ).all()
        energy_records = db.scalars(
            select(EnergyRecord).order_by(EnergyRecord.date.asc(), EnergyRecord.id.asc())
        ).all()
        equipment_service = EquipmentManagementService()
        latest_equipment_records = equipment_service.list_latest(db)
        alerts = equipment_service.analyze_anomalies(db)

        return {
            "production_analysis": self._build_production_analysis(production_records),
            "equipment_analysis": self._build_equipment_analysis(
                latest_equipment_records, alerts
            ),
            "energy_analysis": self._build_energy_analysis(energy_records),
        }

    def generate_business_report(
        self,
        db: Session,
        user_id: int,
        title: str = "制造业生产经营分析报告",
        period_start=None,
        period_end=None,
    ) -> dict:
        """Generate an immutable report using deterministic facts before AI explanation."""
        snapshot = self.build_deterministic_snapshot(db)
        diagnoses = self._build_equipment_diagnoses(db)
        snapshot["equipment_diagnoses"] = diagnoses
        ai_result = self._generate_business_insight(snapshot, diagnoses)
        snapshot["ai_summary"] = {
            "summary": ai_result["summary"],
            "suggestions": ai_result["suggestions"],
            "mode": ai_result["mode"],
        }
        return self.create_snapshot(
            db,
            user_id,
            ManufacturingBusinessReportSnapshotCreate(
                title=title,
                period_start=period_start,
                period_end=period_end,
                risk_level=ai_result["risk_level"],
                ai_mode=ai_result["mode"],
                summary=ai_result["summary"],
                snapshot=snapshot,
                generated_at=datetime.now(timezone.utc),
            ),
        )

    def list_for_user(self, db: Session, user_id: int) -> list[dict]:
        reports = db.scalars(
            select(ManufacturingBusinessReport)
            .where(ManufacturingBusinessReport.user_id == user_id)
            .order_by(
                ManufacturingBusinessReport.generated_at.desc(),
                ManufacturingBusinessReport.id.desc(),
            )
        ).all()
        return [self.serialize(report) for report in reports]

    def build_export(self, report: dict, report_format: str) -> bytes:
        """Render one immutable report dictionary without querying operational data."""
        if report_format == "excel":
            return self._build_excel_export(report)
        if report_format == "word":
            return self._build_word_export(report)
        if report_format == "pdf":
            return self._build_pdf_export(report)
        raise ValueError("报告类型不存在")

    def create_snapshot(
        self,
        db: Session,
        user_id: int,
        payload: ManufacturingBusinessReportSnapshotCreate,
    ) -> dict:
        values = payload.model_dump()
        values["snapshot"] = json.loads(
            json.dumps(values["snapshot"], ensure_ascii=False, default=str)
        )
        report = ManufacturingBusinessReport(user_id=user_id, **values)
        db.add(report)
        db.commit()
        db.refresh(report)
        return self.serialize(report)

    def get_detail(
        self,
        db: Session,
        user_id: int,
        report_id: int,
    ) -> dict | None:
        report = db.scalars(
            select(ManufacturingBusinessReport).where(
                ManufacturingBusinessReport.id == report_id,
                ManufacturingBusinessReport.user_id == user_id,
            )
        ).first()
        return self.serialize(report) if report else None

    @staticmethod
    def _build_equipment_diagnoses(db: Session) -> list[dict]:
        equipment_service = EquipmentManagementService()
        diagnosis_service = EquipmentDiagnosisService()
        diagnoses: list[dict] = []
        for equipment in equipment_service.list_latest(db):
            diagnosis = diagnosis_service.diagnose(db, equipment["equipment_name"])
            if diagnosis is not None:
                diagnoses.append(diagnosis)
        return diagnoses

    def _generate_business_insight(self, snapshot: dict, diagnoses: list[dict]) -> dict:
        fallback = self._build_rule_business_insight(snapshot, diagnoses)
        settings = get_settings()
        if settings.llm_provider != "deepseek" or not settings.llm_api_key:
            return fallback

        prompt = (
            "你是一名制造业生产经营分析师。只能依据 Python 已计算的确定性指标和设备诊断结果，"
            "生成中文经营报告总结。不得重新计算、补充或猜测任何数字，不得虚构产量、能耗、"
            "设备状态、故障、趋势、因果关系或未提供的数据。所有数字只能来自 deterministic_snapshot。"
            "设备诊断只可作为风险解释依据。返回 JSON：summary, risk_level, suggestions。"
            "risk_level 只能是 高风险、中风险、正常；suggestions 必须是字符串数组。\n"
            f"deterministic_snapshot：{snapshot}\n"
            f"equipment_diagnoses：{self._diagnosis_context(diagnoses)}"
        )
        try:
            generated = AIAnalysisService._request_deepseek_json(prompt)
        except Exception:
            return fallback

        risk_level = generated.get("risk_level")
        return {
            "summary": generated.get("summary")
            if isinstance(generated.get("summary"), str) and generated["summary"].strip()
            else fallback["summary"],
            "risk_level": risk_level
            if risk_level in {"高风险", "中风险", "正常"}
            else fallback["risk_level"],
            "suggestions": self._string_list(generated.get("suggestions"), fallback["suggestions"]),
            "mode": "deepseek",
        }

    @staticmethod
    def _build_rule_business_insight(snapshot: dict, diagnoses: list[dict]) -> dict:
        production = snapshot["production_analysis"]
        equipment = snapshot["equipment_analysis"]
        energy = snapshot["energy_analysis"]
        high_risk = any(item["risk_level"] == "高风险" for item in diagnoses)
        medium_risk = any(item["risk_level"] == "中风险" for item in diagnoses)
        completion_rate = production["completion_rate"]

        if high_risk:
            risk_level = "高风险"
        elif medium_risk or (completion_rate is not None and completion_rate < 90):
            risk_level = "中风险"
        else:
            risk_level = "正常"

        summary_parts = [
            f"本期熟料总产量 {production['clinker_output_total']}，水泥总产量 {production['cement_output_total']}。"
        ]
        if completion_rate is not None:
            summary_parts.append(f"水泥产量完成率为 {completion_rate}%。")
        if energy["average_unit_energy_consumption"] is not None:
            summary_parts.append(
                f"平均单位能耗为 {energy['average_unit_energy_consumption']}，总电耗为 {energy['electricity_consumption_total']}。"
            )
        summary_parts.append(
            f"当前纳入分析的设备 {equipment['equipment_count']} 台，异常设备 {equipment['abnormal_equipment_count']} 台。"
        )

        suggestions: list[str] = []
        if high_risk:
            suggestions.append("优先处理高风险设备告警，完成点检与维护闭环后再评估运行风险。")
        if completion_rate is not None and completion_rate < 90:
            suggestions.append("围绕计划差额复盘生产线运行时长、停机原因与产能安排。")
        if energy["average_unit_energy_consumption"] is not None:
            suggestions.append("持续按日跟踪单位能耗趋势，并结合生产线产量同步复盘能耗变化。")
        if not suggestions:
            suggestions.append("保持生产、设备和能源数据的日度记录，持续跟踪经营指标变化。")

        return {
            "summary": "".join(summary_parts),
            "risk_level": risk_level,
            "suggestions": suggestions,
            "mode": "rule_based",
        }

    @staticmethod
    def _diagnosis_context(diagnoses: list[dict]) -> list[dict]:
        return [
            {
                "equipment_name": diagnosis["equipment_name"],
                "risk_level": diagnosis["risk_level"],
                "problem_analysis": diagnosis["problem_analysis"],
                "possible_causes": diagnosis["possible_causes"],
                "suggestions": diagnosis["suggestions"],
            }
            for diagnosis in diagnoses
        ]

    @staticmethod
    def _string_list(value: object, fallback: list[str]) -> list[str]:
        if isinstance(value, list) and all(isinstance(item, str) for item in value):
            return value
        return fallback

    @staticmethod
    def _analysis_sections(report: dict) -> tuple[dict, dict, dict, dict, list[dict]]:
        snapshot = report.get("snapshot") or {}
        return (
            snapshot.get("production_analysis") or {},
            snapshot.get("equipment_analysis") or {},
            snapshot.get("energy_analysis") or {},
            snapshot.get("ai_summary") or {},
            snapshot.get("equipment_diagnoses") or [],
        )

    def _build_excel_export(self, report: dict) -> bytes:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill

        production, equipment, energy, ai_summary, diagnoses = self._analysis_sections(report)
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "生产经营报告"
        sheet.append([report["title"]])
        sheet.merge_cells("A1:B1")
        sheet["A1"].font = Font(bold=True, size=16, color="FFFFFF")
        sheet["A1"].fill = PatternFill("solid", fgColor="1D4ED8")
        sheet.append(["生成时间", report["generated_at"]])
        sheet.append(["AI 总结", report["summary"]])
        sheet.append(["熟料总产量", production.get("clinker_output_total")])
        sheet.append(["水泥总产量", production.get("cement_output_total")])
        sheet.append(["计划总产量", production.get("planned_output_total")])
        sheet.append(["生产完成率", production.get("completion_rate")])
        sheet.append(["设备数量", equipment.get("equipment_count")])
        sheet.append(["设备运行率", equipment.get("running_rate")])
        sheet.append(["故障数量", equipment.get("fault_count")])
        sheet.append(["异常设备数量", equipment.get("abnormal_equipment_count")])
        sheet.append(["平均单位能耗", energy.get("average_unit_energy_consumption")])
        sheet.append(["总电耗", energy.get("electricity_consumption_total")])
        sheet.append(["总煤耗", energy.get("coal_consumption_total")])
        sheet.append([])
        sheet.append(["AI 建议"])
        for suggestion in ai_summary.get("suggestions") or []:
            sheet.append([suggestion])
        if production.get("production_line_comparison"):
            sheet.append([])
            sheet.append(["生产线产量对比"])
            sheet.append(["生产线", "熟料产量", "水泥产量", "计划产量", "完成率"])
            for item in production["production_line_comparison"]:
                sheet.append([
                    item.get("production_line"), item.get("clinker_output"), item.get("cement_output"),
                    item.get("planned_output"), item.get("completion_rate"),
                ])
        if diagnoses:
            sheet.append([])
            sheet.append(["设备诊断"])
            sheet.append(["设备", "风险等级", "问题分析"])
            for diagnosis in diagnoses:
                sheet.append([diagnosis.get("equipment_name"), diagnosis.get("risk_level"), diagnosis.get("problem_analysis")])
        sheet.column_dimensions["A"].width = 26
        sheet.column_dimensions["B"].width = 74
        output = BytesIO()
        workbook.save(output)
        return output.getvalue()

    def _build_word_export(self, report: dict) -> bytes:
        from docx import Document

        production, equipment, energy, ai_summary, diagnoses = self._analysis_sections(report)
        doc = Document()
        doc.add_heading(report["title"], 0)
        doc.add_paragraph(f"生成时间：{report['generated_at']}")
        doc.add_heading("AI 总结", 1)
        doc.add_paragraph(report["summary"])
        doc.add_heading("生产分析", 1)
        self._add_word_kv_table(doc, [
            ("熟料总产量", production.get("clinker_output_total")),
            ("水泥总产量", production.get("cement_output_total")),
            ("计划总产量", production.get("planned_output_total")),
            ("完成率", production.get("completion_rate")),
        ])
        doc.add_heading("设备分析", 1)
        self._add_word_kv_table(doc, [
            ("设备数量", equipment.get("equipment_count")),
            ("运行率", equipment.get("running_rate")),
            ("故障数量", equipment.get("fault_count")),
            ("异常设备数量", equipment.get("abnormal_equipment_count")),
        ])
        doc.add_heading("能源分析", 1)
        self._add_word_kv_table(doc, [
            ("平均单位能耗", energy.get("average_unit_energy_consumption")),
            ("总电耗", energy.get("electricity_consumption_total")),
            ("总煤耗", energy.get("coal_consumption_total")),
        ])
        doc.add_heading("经营建议", 1)
        for suggestion in ai_summary.get("suggestions") or []:
            doc.add_paragraph(suggestion, style="List Bullet")
        if diagnoses:
            doc.add_heading("设备诊断", 1)
            for diagnosis in diagnoses:
                doc.add_paragraph(
                    f"{diagnosis.get('equipment_name')}（{diagnosis.get('risk_level')}）：{diagnosis.get('problem_analysis')}"
                )
        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    @staticmethod
    def _add_word_kv_table(doc, rows: list[tuple[str, object]]) -> None:
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = "--" if value is None else str(value)

    def _build_pdf_export(self, report: dict) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen.canvas import Canvas

        production, equipment, energy, ai_summary, diagnoses = self._analysis_sections(report)
        output = BytesIO()
        canvas = Canvas(output, pagesize=A4)
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        canvas.setFont("STSong-Light", 16)
        canvas.drawString(50, 800, report["title"])
        canvas.setFont("STSong-Light", 10)
        lines = [
            f"生成时间：{report['generated_at']}",
            f"风险等级：{report['risk_level']}",
            f"AI 总结：{report['summary']}",
            f"熟料总产量：{production.get('clinker_output_total')}",
            f"水泥总产量：{production.get('cement_output_total')}",
            f"计划总产量：{production.get('planned_output_total')}",
            f"完成率：{production.get('completion_rate')}",
            f"设备数量：{equipment.get('equipment_count')}，异常设备数量：{equipment.get('abnormal_equipment_count')}",
            f"平均单位能耗：{energy.get('average_unit_energy_consumption')}，总电耗：{energy.get('electricity_consumption_total')}",
            *[f"建议：{item}" for item in ai_summary.get("suggestions") or []],
            *[f"设备诊断：{item.get('equipment_name')} {item.get('risk_level')} {item.get('problem_analysis')}" for item in diagnoses],
        ]
        y_position = 770
        for line in lines:
            canvas.drawString(50, y_position, line[:58])
            y_position -= 24
            if y_position < 60:
                canvas.showPage()
                canvas.setFont("STSong-Light", 10)
                y_position = 800
        canvas.save()
        return output.getvalue()

    @staticmethod
    def _build_production_analysis(records: list[ProductionRecord]) -> dict:
        totals_by_line: dict[str, dict[str, float]] = defaultdict(
            lambda: {"clinker_output": 0.0, "cement_output": 0.0, "planned_output": 0.0}
        )
        clinker_output_total = 0.0
        cement_output_total = 0.0
        planned_output_total = 0.0

        for record in records:
            clinker_output = float(record.clinker_output)
            cement_output = float(record.cement_output)
            planned_output = float(record.planned_output)
            clinker_output_total += clinker_output
            cement_output_total += cement_output
            planned_output_total += planned_output
            line_totals = totals_by_line[record.production_line]
            line_totals["clinker_output"] += clinker_output
            line_totals["cement_output"] += cement_output
            line_totals["planned_output"] += planned_output

        return {
            "clinker_output_total": round(clinker_output_total, 2),
            "cement_output_total": round(cement_output_total, 2),
            "planned_output_total": round(planned_output_total, 2),
            "completion_rate": ManufacturingBusinessReportService._completion_rate(
                cement_output_total, planned_output_total
            ),
            "production_line_comparison": [
                {
                    "production_line": production_line,
                    "clinker_output": round(values["clinker_output"], 2),
                    "cement_output": round(values["cement_output"], 2),
                    "planned_output": round(values["planned_output"], 2),
                    "completion_rate": ManufacturingBusinessReportService._completion_rate(
                        values["cement_output"], values["planned_output"]
                    ),
                }
                for production_line, values in sorted(totals_by_line.items())
            ],
        }

    @staticmethod
    def _build_equipment_analysis(
        latest_records: list[dict], alerts: list[dict]
    ) -> dict:
        equipment_count = len(latest_records)
        running_hours_total = sum(float(record["running_hours"]) for record in latest_records)
        abnormal_equipment_names = {alert["equipment_name"] for alert in alerts}

        return {
            "equipment_count": equipment_count,
            "running_rate": (
                round(running_hours_total / (equipment_count * 24) * 100, 2)
                if equipment_count
                else None
            ),
            "fault_count": sum(int(record["fault_count"]) for record in latest_records),
            "abnormal_equipment_count": len(abnormal_equipment_names),
        }

    @staticmethod
    def _build_energy_analysis(records: list[EnergyRecord]) -> dict:
        totals_by_date: dict[str, dict[str, float]] = defaultdict(
            lambda: {
                "electricity_consumption": 0.0,
                "coal_consumption": 0.0,
                "unit_energy_consumption_total": 0.0,
                "record_count": 0.0,
            }
        )
        electricity_consumption_total = 0.0
        coal_consumption_total = 0.0
        unit_energy_consumption_total = 0.0

        for record in records:
            electricity_consumption = float(record.electricity_consumption)
            coal_consumption = float(record.coal_consumption)
            unit_energy_consumption = float(record.unit_energy_consumption)
            electricity_consumption_total += electricity_consumption
            coal_consumption_total += coal_consumption
            unit_energy_consumption_total += unit_energy_consumption
            date_totals = totals_by_date[record.date.isoformat()]
            date_totals["electricity_consumption"] += electricity_consumption
            date_totals["coal_consumption"] += coal_consumption
            date_totals["unit_energy_consumption_total"] += unit_energy_consumption
            date_totals["record_count"] += 1

        return {
            "average_unit_energy_consumption": (
                round(unit_energy_consumption_total / len(records), 2) if records else None
            ),
            "electricity_consumption_total": round(electricity_consumption_total, 2),
            "coal_consumption_total": round(coal_consumption_total, 2),
            "energy_trend": [
                {
                    "date": record_date,
                    "electricity_consumption": round(values["electricity_consumption"], 2),
                    "coal_consumption": round(values["coal_consumption"], 2),
                    "average_unit_energy_consumption": round(
                        values["unit_energy_consumption_total"] / values["record_count"], 2
                    ),
                }
                for record_date, values in sorted(totals_by_date.items())
            ],
        }

    @staticmethod
    def _completion_rate(actual_output: float, planned_output: float) -> float | None:
        return round(actual_output / planned_output * 100, 2) if planned_output else None

    @staticmethod
    def serialize(report: ManufacturingBusinessReport) -> dict:
        return {
            "id": report.id,
            "title": report.title,
            "period_start": report.period_start.isoformat() if report.period_start else None,
            "period_end": report.period_end.isoformat() if report.period_end else None,
            "risk_level": report.risk_level,
            "ai_mode": report.ai_mode,
            "summary": report.summary,
            "snapshot": report.snapshot,
            "generated_at": report.generated_at.isoformat(),
            "created_at": report.created_at.isoformat(),
        }
