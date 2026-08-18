from io import BytesIO
from pathlib import Path

from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Font, PatternFill

from app.models.dataset import Dataset
from app.services.ai_analysis_service import AIAnalysisService


class ReportService:
    def __init__(self) -> None:
        self.analysis_service = AIAnalysisService()

    def build_excel(self, db, dataset: Dataset) -> bytes:
        analysis = self.analysis_service.generate_report(db, dataset)
        metrics = analysis["metrics"]
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "业务分析报告"
        sheet.append(["AI 智能数据分析助手 - 业务分析报告"])
        sheet.merge_cells("A1:B1")
        sheet["A1"].font = Font(bold=True, size=16, color="FFFFFF")
        sheet["A1"].fill = PatternFill("solid", fgColor="1D4ED8")
        sheet.append(["数据集", dataset.original_filename])
        sheet.append(["有效记录", metrics["total_rows"]])
        overview_rows = self._overview_rows(metrics)
        for label, value in overview_rows:
            sheet.append([label, value])
        skipped_analyses = self._skipped_analysis_rows(metrics)
        if skipped_analyses:
            sheet.append(["本次未分析指标", "；".join(skipped_analyses)])
        for title, headers, rows in [
            *self._order_tables(metrics),
            *self._student_score_tables(metrics),
            *self._inventory_tables(metrics),
        ]:
            sheet.append([])
            sheet.append([title])
            sheet.append(headers)
            for row in rows:
                sheet.append(row)
        sheet.append([])
        sheet.append(["AI 分析模式", analysis["mode"]])
        sheet.append(["数据摘要", analysis["summary"]])
        sheet.append(["异常发现", "；".join(analysis["anomalies"])])
        sheet.append(["业务问题", "；".join(analysis["business_problems"])])
        sheet.append(["优化建议", "；".join(analysis["recommendations"])])
        top_regions = [] if self._is_student_score(metrics) or self._is_inventory(metrics) else metrics.get("top_regions", [])
        if top_regions:
            sheet["D1"] = "区域"
            sheet["E1"] = "已验证销售额"
            for row, region in enumerate(top_regions, start=2):
                sheet.cell(row=row, column=4, value=region["name"])
                sheet.cell(row=row, column=5, value=region["value"])
            chart = BarChart()
            chart.title = "区域已验证销售额 TOP"
            chart.y_axis.title = "已验证销售额"
            chart.x_axis.title = "区域"
            chart.add_data(Reference(sheet, min_col=5, min_row=1, max_row=len(top_regions) + 1), titles_from_data=True)
            chart.set_categories(Reference(sheet, min_col=4, min_row=2, max_row=len(top_regions) + 1))
            chart.height = 8
            chart.width = 15
            sheet.add_chart(chart, "D4")
        sheet.column_dimensions["A"].width = 18
        sheet.column_dimensions["B"].width = 90
        for cell in sheet["A"]:
            cell.font = Font(bold=True)
        output = BytesIO()
        workbook.save(output)
        return output.getvalue()

    def _analysis(self, db, dataset: Dataset) -> dict:
        return self.analysis_service.generate_report(db, dataset)

    @staticmethod
    def _overview_rows(metrics: dict) -> list[tuple[str, str | float | int | None]]:
        """Build a shared business overview for Excel, Word, and PDF reports."""
        if ReportService._is_student_score(metrics):
            student_analysis = ReportService._student_score_analysis(metrics)
            score_summary = student_analysis.get("score_summary")
            rows: list[tuple[str, str | float | int | None]] = []
            if student_analysis.get("student_count") is not None:
                rows.append(("学生数量", student_analysis["student_count"]))
            if isinstance(score_summary, dict):
                rows.extend(
                    [
                        ("有效成绩数量", score_summary.get("count")),
                        ("平均分", score_summary.get("average")),
                        ("中位数", score_summary.get("median")),
                        ("最高分", score_summary.get("maximum")),
                        ("最低分", score_summary.get("minimum")),
                    ]
                )
            return rows
        if ReportService._is_inventory(metrics):
            inventory_analysis = ReportService._inventory_analysis(metrics)
            stock_summary = inventory_analysis.get("stock_summary")
            inventory_value = inventory_analysis.get("inventory_value")
            rows: list[tuple[str, str | float | int | None]] = []
            if inventory_analysis.get("inventory_count") is not None:
                rows.append(("商品数量", inventory_analysis["inventory_count"]))
            if isinstance(stock_summary, dict):
                rows.extend(
                    [
                        ("有效库存记录数", stock_summary.get("count")),
                        ("库存总量", stock_summary.get("total")),
                        ("平均库存", stock_summary.get("average")),
                        ("最大库存", stock_summary.get("maximum")),
                        ("最小库存", stock_summary.get("minimum")),
                        ("库存中位数", stock_summary.get("median")),
                    ]
                )
            if isinstance(inventory_value, dict):
                rows.extend(
                    [
                        ("库存价值总计", inventory_value.get("total")),
                        ("平均库存价值", inventory_value.get("average")),
                    ]
                )
            return rows
        order_analysis = ReportService._order_analysis(metrics)
        if order_analysis:
            overview = order_analysis.get("overview") or {}
            rows: list[tuple[str, str | float | int | None]] = []
            for key, label in (
                ("record_count", "记录数"),
                ("order_count", "订单数量"),
                ("verified_sales_total", "已验证销售额总计"),
                ("verified_order_count", "已验证订单数"),
                ("average_verified_order_value", "已验证平均客单价"),
                ("maximum_verified_order_value", "最大已验证订单金额"),
                ("minimum_verified_order_value", "最小已验证订单金额"),
                ("median_verified_order_value", "已验证订单金额中位数"),
                ("unverified_order_count", "未验证金额订单数"),
                ("unverified_amount_total", "未验证原始金额总计"),
                ("amount_comparable_count", "可比较金额订单数"),
                ("amount_mismatch_count", "金额不一致记录数"),
                ("amount_mismatch_rate", "金额不一致率"),
                ("completed_sales_amount", "已完成订单已验证金额"),
                ("cancelled_order_amount", "已取消订单已验证金额"),
                ("refund_related_amount", "退款相关订单已验证金额"),
            ):
                if overview.get(key) is not None:
                    rows.append((label, overview[key]))
            rows.append(("统计口径", "销售金额优先采用单价 × 数量 × 有效折扣计算得到的可验证金额；仅存在原始订单金额且无法验证的记录不计入已验证销售额。"))
            return rows
        sales = metrics.get("sales_amount")
        highest = metrics.get("highest_sales_region")
        lowest = metrics.get("lowest_sales_region")
        ranking = metrics.get("region_ranking", [])[:5]
        products = metrics.get("product_quantity", [])

        def format_region(region: dict | None) -> str:
            if not region:
                return "-"
            return f"{region['name']}\uff08{region['value']:g}\uff09"

        ranking_text = "\uff1b".join(
            f"{index}. {region['name']}\uff1a{region['value']:g}"
            for index, region in enumerate(ranking, start=1)
        ) or "-"
        completion_rate = metrics.get("completion_rate")
        rows: list[tuple[str, str | float | int | None]] = []
        if metrics.get("order_count") is not None:
            rows.append(("\u8ba2\u5355\u6570\u91cf", metrics["order_count"]))
        if isinstance(sales, dict):
            rows.append(("\u9500\u552e\u989d\u603b\u8ba1", sales.get("total")))
            rows.append(("\u5e73\u5747\u9500\u552e\u989d", sales.get("average")))
        if products:
            product_text = "\uff1b".join(
                f"{index}. {product['name']}\uff1a{product['value']:g}"
                for index, product in enumerate(products[:5], start=1)
            )
            rows.append(("\u5546\u54c1\u9500\u91cf\u6392\u540d", product_text))
        if completion_rate is not None:
            rows.append(("\u5b8c\u6210\u7387", f"{completion_rate}%"))
        if ranking:
            rows.extend(
                [
                    ("\u6700\u5927\u9500\u552e\u533a\u57df", format_region(highest)),
                    ("\u6700\u4f4e\u9500\u552e\u533a\u57df", format_region(lowest)),
                    ("\u533a\u57df TOP \u6392\u540d", ranking_text),
                ]
            )
        return rows

    @staticmethod
    def _is_student_score(metrics: dict) -> bool:
        return (metrics.get("selected_module") or {}).get("id") == "student_score"

    @staticmethod
    def _is_inventory(metrics: dict) -> bool:
        return (metrics.get("selected_module") or {}).get("id") == "inventory"

    @staticmethod
    def _student_score_analysis(metrics: dict) -> dict:
        value = metrics.get("student_score_analysis")
        return value if isinstance(value, dict) else {}

    @staticmethod
    def _inventory_analysis(metrics: dict) -> dict:
        value = metrics.get("inventory_analysis")
        return value if isinstance(value, dict) else {}

    @staticmethod
    def _order_analysis(metrics: dict) -> dict:
        value = metrics.get("order_analysis")
        return value if isinstance(value, dict) else {}

    @staticmethod
    def _order_tables(metrics: dict) -> list[tuple[str, list[str], list[list[object]]]]:
        analysis = ReportService._order_analysis(metrics)
        if not analysis:
            return []
        tables: list[tuple[str, list[str], list[list[object]]]] = []
        for key, title, headers, fields in (
            ("product_analysis", "商品销售分析", ["商品", "订单数", "销量", "已验证销售额"], ["name", "order_count", "quantity", "sales_amount"]),
            ("category_analysis", "品类销售分析", ["品类", "订单数", "销量", "已验证销售额", "销售占比"], ["category", "order_count", "quantity", "sales_amount", "sales_share"]),
            ("region_analysis", "地区销售分析", ["地区", "订单数", "销量", "已验证销售额", "已验证平均客单价"], ["name", "region_order_count", "region_quantity", "region_sales", "region_average_order_value"]),
            ("status_analysis", "订单状态分析", ["状态", "订单数", "已验证订单金额", "订单占比"], ["name", "order_count", "sales_amount", "share"]),
            ("payment_method_analysis", "支付方式分析", ["支付方式", "订单数", "已验证订单金额", "订单占比", "销售额占比"], ["name", "order_count", "sales_amount", "order_share", "sales_share"]),
        ):
            values = analysis.get(key) or []
            if values:
                tables.append((title, headers, [[item.get(field) for field in fields] for item in values]))
        customer = analysis.get("customer_analysis")
        if isinstance(customer, dict) and customer.get("top_customers"):
            tables.append((
                "客户复购分析",
                ["客户编号", "订单数", "已验证销售额"],
                [[item.get("customer_id"), item.get("order_count"), item.get("sales_amount")] for item in customer["top_customers"]],
            ))
        time_analysis = analysis.get("time_analysis") or {}
        for key, title, header in (("daily_sales_trend", "日销售趋势", "日期"), ("monthly_sales_trend", "月度销售趋势", "月份"), ("monthly_order_trend", "月度订单趋势", "月份")):
            values = time_analysis.get(key) or []
            if values:
                tables.append((title, [header, "数值"], [[item.get("name"), item.get("value")] for item in values]))
        discount = analysis.get("discount_analysis")
        if isinstance(discount, dict):
            tables.append(("折扣分析", list(discount.keys()), [list(discount.values())]))
        quality = analysis.get("data_quality")
        if isinstance(quality, dict):
            labels = {
                "contact_complete_rate": "联系方式完整率(%)",
                "contact_complete_count": "联系方式完整记录数",
                "phone_invalid_count": "无效联系方式数",
                "email_invalid_count": "无效邮箱数",
                "phone_missing_count": "缺失联系方式数",
                "email_missing_count": "缺失邮箱数",
            }
            rows = [
                [labels.get(key, key), value]
                for key, value in quality.items()
                if key != "missing_value_summary" and not isinstance(value, (dict, list))
            ]
            if quality.get("missing_value_summary"):
                rows.append(["missing_value_summary", "；".join(f"{key}:{value}" for key, value in quality["missing_value_summary"].items())])
            if rows:
                tables.append(("数据质量检查", ["检查项", "结果"], rows))
        return tables

    @staticmethod
    def _student_score_tables(metrics: dict) -> list[tuple[str, list[str], list[list[object]]]]:
        if not ReportService._is_student_score(metrics):
            return []
        analysis = ReportService._student_score_analysis(metrics)
        tables: list[tuple[str, list[str], list[list[object]]]] = []
        for key, title, first_column in (
            ("subject_score", "学科成绩统计", "学科"),
            ("class_score", "班级成绩统计", "班级"),
        ):
            values = analysis.get(key, [])
            if values:
                tables.append(
                    (
                        title,
                        [first_column, "有效成绩数", "平均分", "最高分", "最低分"],
                        [
                            [item["name"], item["count"], item["average"], item["maximum"], item["minimum"]]
                            for item in values
                        ],
                    )
                )
        students = analysis.get("student_score", [])
        if students:
            headers = ["学生编号", "有效成绩数", "平均分", "最高分", "最低分"]
            if any(item.get("student_name") for item in students):
                headers.insert(1, "学生姓名")
            rows: list[list[object]] = []
            for item in students:
                row: list[object] = [item["student_id"]]
                if "学生姓名" in headers:
                    row.append(item.get("student_name", ""))
                rows.append(row + [item["score_count"], item["average"], item["maximum"], item["minimum"]])
            tables.append(("学生成绩汇总", headers, rows))
        trend = analysis.get("exam_trend", [])
        if trend:
            tables.append(
                (
                    "考试趋势",
                    ["考试", "平均分", "有效成绩数"],
                    [[item["name"], item["average"], item["count"]] for item in trend],
                )
            )
        return tables

    @staticmethod
    def _inventory_tables(metrics: dict) -> list[tuple[str, list[str], list[list[object]]]]:
        if not ReportService._is_inventory(metrics):
            return []
        analysis = ReportService._inventory_analysis(metrics)
        tables: list[tuple[str, list[str], list[list[object]]]] = []
        low_stock = analysis.get("low_stock_analysis", [])
        if low_stock:
            tables.append(
                (
                    "低库存明细",
                    ["商品编号", "商品名称", "当前库存", "安全库存", "库存缺口"],
                    [
                        [
                            item.get("product_id", ""),
                            item.get("product_name", ""),
                            item.get("stock_quantity"),
                            item.get("safety_stock"),
                            item.get("shortage"),
                        ]
                        for item in low_stock
                    ],
                )
            )
        for key, title, first_column in (
            ("category_stock", "分类库存统计", "分类"),
            ("warehouse_stock", "仓库库存统计", "仓库"),
            ("supplier_stock", "供应商库存统计", "供应商"),
        ):
            values = analysis.get(key, [])
            if values:
                tables.append(
                    (title, [first_column, "库存数量"], [[item["name"], item["value"]] for item in values])
                )
        inventory_flow = analysis.get("inventory_flow")
        if isinstance(inventory_flow, dict):
            tables.append(
                (
                    "库存流动统计",
                    ["入库总量", "出库总量", "净变化"],
                    [[inventory_flow.get("inbound_total"), inventory_flow.get("outbound_total"), inventory_flow.get("net_change")]],
                )
            )
        inventory_trend = analysis.get("inventory_trend", [])
        if inventory_trend:
            tables.append(
                ("库存趋势", ["日期", "库存总量"], [[item["name"], item["value"]] for item in inventory_trend])
            )
        return tables

    @staticmethod
    def _skipped_analysis_rows(metrics: dict) -> list[str]:
        """Return compact capability notices without treating unavailable data as zero."""
        rows: list[str] = []
        for item in metrics.get("analysis_plan", []):
            if item.get("supported"):
                continue
            missing_fields = item.get("missing_fields", [])
            detail = (
                f"\u7f3a\u5c11 {'\u3001'.join(missing_fields)} \u5b57\u6bb5"
                if missing_fields
                else item.get("reason") or "\u672a\u751f\u6210\u53ef\u7528\u8ba1\u7b97\u7ed3\u679c"
            )
            rows.append(f"{item.get('name', item.get('id', '\u672a\u77e5\u6307\u6807'))}\uff1a{detail}")
        return rows

    @staticmethod
    def _resolve_chart_font_path() -> Path:
        """返回当前运行环境中可供 Matplotlib 使用的中文字体文件。"""
        candidates = (
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
            Path("/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc"),
            Path(r"C:\Windows\Fonts\msyh.ttc"),
            Path(r"C:\Windows\Fonts\simhei.ttf"),
        )
        for font_path in candidates:
            if font_path.is_file():
                return font_path
        raise RuntimeError("未找到可用中文字体，请安装 Noto CJK 或配置运行环境字体")

    @staticmethod
    def _build_region_chart(metrics: dict) -> BytesIO | None:
        regions = metrics.get("top_regions", [])
        if not regions:
            return None

        import matplotlib

        matplotlib.use("Agg")
        from matplotlib import pyplot as plt
        from matplotlib.font_manager import FontProperties

        plt.rcParams["axes.unicode_minus"] = False
        chinese_font = FontProperties(fname=str(ReportService._resolve_chart_font_path()))
        chart, axis = plt.subplots(figsize=(8, 3.8))
        axis.bar(
            [item["name"] for item in regions],
            [item["value"] for item in regions],
            color="#2563EB",
            label="区域已验证销售额",
        )
        axis.set_title("区域已验证销售额 TOP", fontproperties=chinese_font)
        axis.set_xlabel("区域", fontproperties=chinese_font)
        axis.set_ylabel("已验证销售额", fontproperties=chinese_font)
        axis.legend(prop=chinese_font)
        for label in axis.get_xticklabels():
            label.set_fontproperties(chinese_font)
        axis.tick_params(axis="x", rotation=25)
        chart.tight_layout()
        output = BytesIO()
        chart.savefig(output, format="png", dpi=160)
        plt.close(chart)
        output.seek(0)
        return output

    @staticmethod
    def _build_pdf_summary_image(dataset: Dataset, analysis: dict) -> BytesIO:
        """Render PDF text as a CJK-capable image for consistent PDF viewing."""
        import matplotlib

        matplotlib.use("Agg")
        from matplotlib import pyplot as plt
        from matplotlib.font_manager import FontProperties

        chinese_font = FontProperties(fname=str(ReportService._resolve_chart_font_path()))
        from textwrap import wrap

        metrics = analysis["metrics"]
        overview_rows = ReportService._overview_rows(metrics)
        skipped_analyses = ReportService._skipped_analysis_rows(metrics)
        sections = [
            ("AI \u667a\u80fd\u6570\u636e\u5206\u6790\u52a9\u624b - \u4e1a\u52a1\u5206\u6790\u62a5\u544a", 16),
            (f"\u6570\u636e\u96c6\uff1a{dataset.original_filename}    \u6709\u6548\u8bb0\u5f55\uff1a{metrics['total_rows']}", 9),
            *[(f"{label}\uff1a{value}", 8.5) for label, value in overview_rows],
            (f"\u6570\u636e\u6458\u8981\uff1a{analysis['summary']}", 8.5),
            (f"\u5f02\u5e38\u5206\u6790\uff1a{'\uff1b'.join(analysis['anomalies'])}", 8.5),
            (f"\u4e1a\u52a1\u5efa\u8bae\uff1a{'\uff1b'.join(analysis['recommendations'])}", 8.5),
        ]
        if skipped_analyses:
            sections.insert(2, (f"\u672c\u6b21\u672a\u5206\u6790\u6307\u6807\uff1a{'\uff1b'.join(skipped_analyses)}", 8.5))
        if ReportService._is_student_score(metrics):
            sections[0] = ("AI 智能数据分析助手 - 学生成绩分析报告", 16)
            for title, headers, rows in ReportService._student_score_tables(metrics):
                sections.append((title, 10))
                sections.append(("；".join(headers), 8.5))
                sections.extend(
                    ("；".join(str(value) for value in row), 8.5)
                    for row in rows
                )
        if ReportService._is_inventory(metrics):
            sections[0] = ("AI 智能数据分析助手 - 库存分析报告", 16)
            for title, headers, rows in ReportService._inventory_tables(metrics):
                sections.append((title, 10))
                sections.append(("；".join(headers), 8.5))
                sections.extend(
                    ("；".join(str(value) for value in row), 8.5)
                    for row in rows
                )
        for title, headers, rows in ReportService._order_tables(metrics):
            sections.append((title, 10))
            sections.append(("；".join(headers), 8.5))
            sections.extend(("；".join(str(value) for value in row), 8.5) for row in rows)
        figure, axis = plt.subplots(figsize=(8.2, max(5.5, 1.2 + len(sections) * 0.42)))
        axis.axis("off")
        y_position = 0.96
        for content, font_size in sections:
            wrapped_content = "\n".join(wrap(content, width=56))
            axis.text(
                0.02,
                y_position,
                wrapped_content,
                fontproperties=chinese_font,
                fontsize=font_size,
                fontweight="bold" if font_size >= 16 else "normal",
                va="top",
            )
            line_height = 0.065 if font_size >= 16 else 0.042
            y_position -= line_height * (wrapped_content.count("\n") + 1) + 0.012
        figure.tight_layout(pad=0.3)
        output = BytesIO()
        figure.savefig(output, format="png", dpi=160, bbox_inches="tight")
        plt.close(figure)
        output.seek(0)
        return output

    def build_word(self, db, dataset: Dataset) -> bytes:
        from docx import Document
        from docx.shared import Pt
        analysis = self._analysis(db, dataset)
        metrics = analysis["metrics"]
        doc = Document()
        doc.add_heading("AI 智能数据分析助手 - 业务分析报告", 0)
        doc.add_paragraph(f"数据集：{dataset.original_filename}")
        doc.add_heading("数据概览", 1)
        doc.add_paragraph(f"有效记录：{metrics['total_rows']}")
        for label, value in self._overview_rows(metrics):
            doc.add_paragraph(f"{label}：{value}")
        skipped_analyses = self._skipped_analysis_rows(metrics)
        if skipped_analyses:
            doc.add_heading("本次未分析指标", 1)
            for item in skipped_analyses:
                doc.add_paragraph(item, style="List Bullet")
        for title, headers, rows in [
            *self._order_tables(metrics),
            *self._student_score_tables(metrics),
            *self._inventory_tables(metrics),
        ]:
            doc.add_heading(title, 1)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = str(header)
            for row in rows:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = str(value)
        chart = None if self._is_student_score(metrics) or self._is_inventory(metrics) else self._build_region_chart(metrics)
        if chart is not None:
            doc.add_heading("区域已验证销售额图表", 1)
            doc.add_picture(chart)
        for title, content in [("数据摘要", analysis["summary"]), ("异常发现", "；".join(analysis["anomalies"])), ("业务问题", "；".join(analysis["business_problems"])), ("优化建议", "；".join(analysis["recommendations"]))]:
            doc.add_heading(title, 1); doc.add_paragraph(content)
        for style in doc.styles:
            if hasattr(style, "font"):
                style.font.name = "Microsoft YaHei"
                style.font.size = Pt(10)
        output = BytesIO(); doc.save(output); return output.getvalue()

    def build_pdf(self, db, dataset: Dataset) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.utils import ImageReader
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen.canvas import Canvas
        analysis = self._analysis(db, dataset)
        output = BytesIO()
        canvas = Canvas(output, pagesize=A4)
        canvas.setTitle("AI Business Analysis Report")
        summary_image = self._build_pdf_summary_image(dataset, analysis)
        canvas.drawImage(
            ImageReader(summary_image),
            50,
            80,
            width=490,
            height=680,
            preserveAspectRatio=True,
            mask="auto",
        )
        canvas.showPage()
        chart = (
            None
            if self._is_student_score(analysis["metrics"]) or self._is_inventory(analysis["metrics"])
            else self._build_region_chart(analysis["metrics"])
        )
        if chart is not None:
            canvas.drawImage(
                ImageReader(chart),
                50,
                260,
                width=490,
                height=300,
                preserveAspectRatio=True,
                mask="auto",
            )
        canvas.save()
        return output.getvalue()

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        output = BytesIO(); canvas = Canvas(output, pagesize=A4); canvas.setFont("STSong-Light", 16)
        canvas.drawString(50, 800, "AI 智能数据分析助手 - 业务分析报告")
        canvas.setFont("STSong-Light", 10); y = 770
        lines = [f"数据集：{dataset.original_filename}", f"有效记录：{analysis['metrics']['total_rows']}", f"销售额总计：{analysis['metrics']['sales_amount']['total']}", f"数据摘要：{analysis['summary']}", f"异常发现：{'；'.join(analysis['anomalies'])}", f"优化建议：{'；'.join(analysis['recommendations'])}"]
        for line in lines:
            canvas.drawString(50, y, line[:70]); y -= 28
        chart = self._build_region_chart(analysis["metrics"])
        if chart is not None:
            canvas.drawImage(ImageReader(chart), 50, 390, width=490, height=232, preserveAspectRatio=True, mask="auto")
        canvas.save(); return output.getvalue()
